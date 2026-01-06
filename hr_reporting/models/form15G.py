from odoo import models
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

class HrCustomFormFifteenG(models.Model):
    _inherit = "hr.custom.form.form15g"
    
    def action_generate_excel_report(self):
        self.ensure_one()

        doc = Document()
        
       # ===== Common Style =====
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        def heading(text):
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True

        def space():
            doc.add_paragraph("")

        def table(rows, cols, widths=None):
            t = doc.add_table(rows=rows, cols=cols)
            t.style = "Table Grid"
            return t

        # =========================================================
        # HEADER
        # =========================================================
        doc.add_paragraph(f"Date : {self.declaration_date.strftime('%d-%m-%Y') or ''}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading("FORM NO. 15G")
        heading("[See section 197A(1), 197A(1A) and rule 29C]")
        heading("Declaration under section 197A (1) and section 197A(1A) to be made by an individual or a person (not being a company or firm) claiming certain incomes without deduction of tax.")
        space()
        heading("PART - I")
        space()

        # Table
        t1 = table(15, 4)

        t1.cell(0,0).text = "1. Name of Assessee (Declarant)"
        t1.cell(0,1).text = self.assessee_name or ""
        t1.cell(0,2).text = "2. PAN of the Assessee"
        t1.cell(0,3).text = self.assessee_pan or ""

        t1.cell(1,0).text = "3. Status"
        t1.cell(1,1).text = self.assessee_status or ""
        t1.cell(1,2).text = "4. Previous Year"
        t1.cell(1,3).text = str(self.previous_year or "")

        t1.cell(2,0).text = "5. Residential Status"
        t1.cell(2,1).merge(t1.cell(2,3))
        t1.cell(2,1).text = self.residential_status or ""

        t1.cell(3,0).text = "6. Flat/Door/Block No."
        t1.cell(3,1).text = str(self.address_flat or "")
        t1.cell(3,2).text = "7. Name of Premises"
        t1.cell(3,3).text = self.address_premises or ""

        t1.cell(4,0).text = "8. Road / Street / Lane"
        t1.cell(4,1).text = self.address_road or ""
        t1.cell(4,2).text = "9. Area / Locality"
        t1.cell(4,3).text = self.address_area or ""

        t1.cell(5,0).text = "10. Town / City / District"
        t1.cell(5,1).merge(t1.cell(5,3))
        t1.cell(5,1).text = self.address_city or ""
        
        t1.cell(6,0).text = "11. State"
        t1.cell(6,1).text = self.address_state or ""
        t1.cell(6,2).text = "12. PIN"
        t1.cell(6,3).text = str(self.address_pin or "") 
        
        t1.cell(7,0).text = "13. Email"
        t1.cell(7,1).merge(t1.cell(7,3))
        t1.cell(7,1).text = self.contact_email or ""

        t1.cell(8,0).merge(t1.cell(8,2))
        t1.cell(8,0).text = "14. Telephone No. (with STD Code) and Mobile No. "
        t1.cell(8,3).text = str(self.contact_phone or "")
        
        t1.cell(9,0).merge(t1.cell(9,2))
        t1.cell(9, 2).text = (
            f"15 (a) Whether assessed to tax under the Income-tax Act, 1961:\n"
            f"     (b) If yes, latest assessment year for which assessed"
        )
        t1.cell(9,3).text = (
            f"YES {'☑' if self.assessed_to_tax == 'yes' else '☐'}        "
            f"NO {'☑' if self.assessed_to_tax == 'no' else '☐'}\n"
            f"{self.latest_assessment_year or ''}"
            )

        t1.cell(10,0).merge(t1.cell(10,2))
        t1.cell(10,0).text = "16. Estimated income for which this declaration is made"
        t1.cell(10,3).text = str(self.estimated_income or "")

        t1.cell(11,0).merge(t1.cell(11,2))
        t1.cell(11,0).text = "17. Estimated total income of the P.Y. in which income mentioned in column 16 to be included"
        t1.cell(11,3).text = str(self.estimated_total_income or "")

        t1.cell(12,0).merge(t1.cell(12,2))
        t1.cell(12,0).text = "18. Details of Form No. 15G other than this form filed during the previous year, if any"
        t1.cell(12,3).text = str(self.other_form15g_details or "")
        
        t1.cell(13,0).merge(t1.cell(13,1))
        t1.cell(13,0).text = "Total No. of Form No. 15G filed"
        t1.cell(13,2).merge(t1.cell(13,3))
        t1.cell(13,2).text = "Aggregate amount of income for which Form No.15G filed"

        t1.cell(14,0).merge(t1.cell(14,1))
        t1.cell(14,0).text = str(self.other_form15g_count or "")
        t1.cell(14,2).merge(t1.cell(14,3))
        t1.cell(14,2).text = str(self.other_form15g_amount or "")

        # Table 2
        t2 = table(len(self.income_detail_ids)+2, 5) 
          
        t2.cell(0,0).merge(t2.cell(0,4))
        t2.cell(0,0).text = "19. Details of income for which the declaration is filed "

        t2.cell(1,0).text = "Sl. No."
        t2.cell(1,1).text = "Identification number of relevant investment/account, etc."
        t2.cell(1,2).text = "Nature of income"
        t2.cell(1,3).text = "Section under which tax is deductible"
        t2.cell(1,4).text = "Amount of income"
        
        row_index = 2  

        for rec in self.income_detail_ids:
            t2.cell(row_index, 0).text = str(rec.sequence or "")
            t2.cell(row_index, 1).text = rec.investment_identification or ""
            t2.cell(row_index, 2).text = rec.income_nature or ""
            t2.cell(row_index, 3).text = rec.deduction_section or ""
            t2.cell(row_index, 4).text = str(rec.income_amount or "")

            row_index += 1

        space()
        space()
        
        doc.add_paragraph(f"Signature of the Declarant:______________________________________")
        
        space()
        space()
        
        # DECLARATION
        heading("Declaration/Verification")
        doc.add_paragraph(
            "*I/We………………………………do hereby declare that to the best of *my/our knowledge and belief what is "
            "stated above is correct, complete and is truly stated. *I/We declare that the incomes referred to "
            "in this form are not includible in the total income of any other person under sections 60 to 64 of the "
            "Income-tax Act, 1961. *I/We further declare that the tax *on my/our estimated total income including "
            "*income/incomes referred to in column 16 *and aggregate amount of *income/incomes referred to in column "
            "18 computed in accordance with the provisions of the Income-tax Act, 1961, for the previous year ending "
            "on .................... relevant to the assessment year ..................will be nil. *I/We also declare"
            "that *my/our *income/incomes referred to in column 16 *and the aggregate amount of *income/incomes referred "
            "to in column 18 for the previous year ending on .................... relevant to the assessment year "
            ".................. will not exceed the maximum amount which is not chargeable to income-tax."
        )

        space()
        doc.add_paragraph("Place: __________________________                       Date : _________________________\n")
        doc.add_paragraph("Signature of Declarant: ________________________")

        space()
        heading("PART - II")
        heading("To be filled by the person responsible for paying the income referred to in column 16 of Part I")

        # Table 3
        t3 = table(10, 2)

        t3.cell(0,0).text = "1. Name of the person responsible for paying"
        t3.cell(0,1).text = self.payer_name or ""

        t3.cell(1,0).text = "2. Unique Identification No."
        t3.cell(1,1).text = self.payer_uid or ""
        
        t3.cell(2,0).text = "3. PAN of the person responsible for paying "
        t3.cell(2,1).text = self.payer_pan or ""
        
        t3.cell(3,0).text = "4. Complete Address "
        t3.cell(3,1).text = self.payer_address or ""
        
        t3.cell(4,0).text = "5. TAN of the person responsible for paying"
        t3.cell(4,1).text = self.payer_tan or ""
        
        t3.cell(5,0).text = "6. Email "
        t3.cell(5,1).text = self.payer_email or ""
        
        t3.cell(6,0).text = "7. Telephone No. (with STD Code) and Mobile No."
        t3.cell(6,1).text = str(self.payer_phone or "")
        
        t3.cell(7,0).text = "8. Amount of income paid"
        t3.cell(7,1).text = str(self.payer_income_amount or "")
        
        t3.cell(8,0).text = "9. Date on which Declaration is received (DD/MM/YYYY) "
        t3.cell(8,1).text = str(self.declaration_received_date or "")
        
        t3.cell(9,0).text = "10. Date on which the income has been paid/credited (DD/MM/YYYY) "
        t3.cell(9,1).text = str(self.income_paid_date or "")

        space()
        doc.add_paragraph("Date: _________________________        Place: __________________________")
        doc.add_paragraph("_____________________________________________________________________________\n"
                          "Signature of the person responsible for paying the income referred to in column 16 of Part I")
        
        
        doc.add_paragraph(
            "*Delete whichever is not applicable."
            "1As per provisions of section 206AA(2), the declaration under section 197A(1) or 197A(1A) shall be invalid if the" 
            "declarant fails to furnish his valid Permanent Account Number (PAN).\n"
            "2Declaration can be furnished by an individual under section 197A(1) and a person (other than a company or a firm)"
            "under section 197A(1A).\n" 
            "3The financial year to which the income pertains.\n"
            "4Please mention the residential status as per the provisions of section 6 of the Income-tax Act, 1961.\n" 
            "5 Please mention “Yes” if assessed to tax under the provisions of Income-tax Act, 1961 for any of the assessment"
            "year out of six assessment years preceding the year in which the declaration is filed.\n" 
            "6Please mention the amount of estimated total income of the previous year for which the declaration is filed" 
            "including the amount of income for which this declaration is made.\n" 
            "7In case any declaration(s) in Form No. 15G is filed before filing this declaration during the previous year, mention" 
            "the total number of such Form No. 15G filed along with the aggregate amount of income for which said" 
            "declaration(s) have been filed.\n" 
            "8Mention the distinctive number of shares, account number of term deposit, recurring deposit, National Savings"
            "Schemes, life insurance policy number, employee code, etc.\n"  
            "9Indicate the capacity in which the declaration is furnished on behalf of a HUF, AOP, etc.\n" 
            "10Before signing the declaration/verification, the declarant should satisfy himself that the information furnished in" 
            "this form is true, correct and complete in all respects. Any person making a false statement in the declaration shall"
            "be liable to prosecution under section 277 of the Income-tax Act, 1961 and on conviction be punishable-\n" 
            "(i) in a case where tax sought to be evaded exceeds twenty-five lakh rupees, with rigorous imprisonment"
            "which shall not be less than six months but which may extend to seven years and with fine;\n"
            "(ii) in any other case, with rigorous imprisonment which shall not be less than three months but which may" 
            "extend to two years and with fine.\n"
            "11The person responsible for paying the income referred to in column 16 of Part I shall allot a unique identification"
            "number to all the Form No. 15G received by him during a quarter of the financial year and report this reference"
            "number along with the particulars prescribed in rule 31A(4)(vii) of the Income-tax Rules, 1962 in the TDS statement" 
            "furnished for the same quarter.  In case the person has also received Form No.15H during the same quarter, please"
            "allot separate series of serial number for Form No.15G and Form No.15H.\n"
            "12The person responsible for paying the income referred to in column 16 of Part I shall not accept the declaration"
            "where the amount of income of the nature referred to in sub-section (1) or sub-section (1A) of section 197A or the"
            "aggregate of the amounts of such income credited or paid or likely to be credited or paid during the previous year in" 
            "which such income is to be included exceeds the maximum amount which is not chargeable to tax. For deciding the"
            "eligibility, he is required to verify income or the aggregate amount of incomes, as the case may be, reported by the" 
            "declarant in columns 16 and 18.;"
        )

        # ============ SAVE ==============
      
        buffer = io.BytesIO()
        doc.save(buffer)
        
        filename = (f"Form 15G ({self.name or ''}).docx")


        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(buffer.getvalue()),
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'res_model': self._name,
            'res_id': self.id,
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }
