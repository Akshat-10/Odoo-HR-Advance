# -*- coding: utf-8 -*-
from odoo import models, fields
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import base64
from datetime import date


class HrCustomFormStaffLoan(models.Model):
    _name = "hr.custom.form.staff_loan"
    _description = "Staff Loan Application"
    _inherit = "hr.custom.form.base"

    _sequence_code = "hr.custom.form.staff_loan"

    word_file = fields.Binary(string="Word File")
    word_filename = fields.Char(string="Word Filename")


    designation = fields.Char(string="Designation")

    loan_amount = fields.Float(string="Loan Amount")
    loan_purpose = fields.Text(string="Purpose of Loan")
    deduction_per_month = fields.Float(string="Deduction per Month from Salary")
    remarks = fields.Text(string="Guarantor Details Remark")

    # ✅ NEW FIELDS (Bank Details)
    bank_name_as_per_bank = fields.Char(string="Name as per Bank")
    bank_name = fields.Char(string="Bank Name")
    bank_account_number = fields.Char(string="Account Number")
    bank_ifsc_code = fields.Char(string="IFSC Code")

    # ✅ Approval of Loan Sanction Committee (NO ERROR VERSION)
    committee_member_1_id = fields.Many2one(
        "hr.employee",
        string="Name",
        domain="[('company_id', '=', company_id)]",
    )
    committee_signature_1 = fields.Char(string="Signature")

    committee_member_2_id = fields.Many2one(
        "hr.employee",
        string="Name",
        domain="[('company_id', '=', company_id)]",
    )
    committee_signature_2 = fields.Char(string="Signature")

    # Display-only fields (Name without code)
    committee_member_1_name = fields.Char(
        string="Name",
        related="committee_member_1_id.name",
        store=False,
    )

    committee_member_2_name = fields.Char(
        string="Name",
        related="committee_member_2_id.name",
        store=False,
    )

    def action_download_staff_loan_word(self):
        self.ensure_one()

        doc = Document()

        # ===== FIXED MARGINS (INT ONLY) =====
        section = doc.sections[0]
        section.top_margin = 360000  # 0.4 inch
        section.bottom_margin = 360000
        section.left_margin = 360000
        section.right_margin = 360000

        def compact(p, size=11, bold=False, underline=False, align=None):
            if align:
                p.alignment = align
            run = p.runs[0]
            run.font.size = Pt(size)
            run.bold = bold
            run.underline = underline
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

        # ================= COMPANY NAME =================
        p = doc.add_paragraph("GERMAN TMX PVT LTD")
        compact(p, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # ================= TITLE =================
        p = doc.add_paragraph("Staff Loan Application")
        compact(p, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # ================= DATE =================
        p = doc.add_paragraph(f"Date :- {date.today().strftime('%d/%m/%Y')}")
        compact(p)

        # ================= BODY =================
        doc.add_paragraph(
            f"Kindly Sanction an Amount of {int(self.loan_amount or 0):,}/- "
            f"as advance required by me for {self.loan_purpose or ''}, "
            f"kindly deduct above sanction amount from my salary "
            f"@RS {int(self.deduction_per_month or 0):,}/- Per Month."
        )
        compact(p)
        # doc.add_paragraph("")
        doc.add_paragraph(f"Name of Applicant: {self.employee_id.name or ''}")
        doc.add_paragraph(f"Code No: {self.employee_id.employee_code or ''}")
        doc.add_paragraph(f"Department: {self.employee_id.department_id.name or ''}")
        doc.add_paragraph(f"Designation: {self.designation or ''}")
        doc.add_paragraph("Signature:____________________________")

        # ================= DECLARATION =================
        p = doc.add_paragraph("\nDeclaration")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.underline = True
        run.font.size = Pt(15)

        doc.add_paragraph(
            f"We named below hereby that we stand as sureties for the amount to be {self.employee_id.name or ''} "
            "and GERMAN TMX PVT LTD  is authorized for recovery amount through guarantor, "
            "if above named applicant fails to repay the issues amount."
        )

        # ================= BANK DETAILS =================
        p = doc.add_paragraph("\nBank Details:")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.underline = True
        run.font.size = Pt(15)

        doc.add_paragraph(f"Name as per bank: {self.bank_name_as_per_bank or ''}")
        doc.add_paragraph(f"Bank Name: {self.bank_name or ''}")
        doc.add_paragraph(
            f"Account Number: {self.bank_account_number or 'As Per HR Record'}"
        )
        doc.add_paragraph(f"IFSC Code: {self.bank_ifsc_code or ''}")

        # ================= GUARANTOR =================
        p = doc.add_paragraph("\nGuarantor Details:")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.underline = True
        run.font.size = Pt(15)

        doc.add_paragraph(self.remarks or "NA")

        # ================= COMMITTEE =================
        p = doc.add_paragraph("Approval of Loan Sanction Committee")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.underline = True
        run.font.size = Pt(15)

        # Table: 4 columns, 3 rows
        table = doc.add_table(rows=3, cols=4)
        table.autofit = True

        # ---------- ROW 0 : HEADER ----------
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Signature"
        table.cell(0, 2).text = "Name"
        table.cell(0, 3).text = "Signature"

        # ---------- ROW 1 : COMMITTEE MEMBER NAMES (FROM IDs) ----------
        table.cell(1, 0).text = self.committee_member_1_id.name if self.committee_member_1_id else ""
        table.cell(1, 2).text = self.committee_member_2_id.name if self.committee_member_2_id else ""

        # ---------- ROW 2 : DESIGNATION ----------
        table.cell(2, 0).text = "()"
        table.cell(2, 2).text = "()"

        # ================= SAVE =================
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        self.word_file = base64.b64encode(buffer.read())
        employee_name = (self.employee_id.name or "").replace(" ", "_")
        self.word_filename = f"Staff_Loan_Application_{employee_name}.docx"

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/?model={self._name}&id={self.id}"
                   f"&field=word_file&filename_field=word_filename&download=true",
            "target": "self",
        }


