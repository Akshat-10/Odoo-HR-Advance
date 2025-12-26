from odoo import models
from docx import Document
from docx.shared import Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import os

class Form15G(models.Model):
    _inherit = "hr.applicant"
    
    def action_generate_excel_report(self):
        self.ensure_one()

        doc = Document()
        
        # ========== Helper Function ==========    
        # Paragraph text (bold / unbold / align)
        def ptext(text, bold=False, align=None):
            p = doc.add_paragraph()
            run = p.add_run(text or "")
            run.bold = bold
            if align:
                p.alignment = align
            return p


        # Table create
        def table(rows, cols, style="Table Grid"):
            t = doc.add_table(rows=rows, cols=cols)
            t.style = style
            return t


        # Table cell text (bold / unbold)
        def ctext(cell, text, bold=False, align=None):
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text or "")
            run.bold = bold
            if align:
                p.alignment = align
        
        # =====================================
       
        t = doc.add_table(2, 3)
        
        t.cell(0, 0).merge(t.cell(0, 1))
        t.cell(0,0).text = f"DEPARTMENT : {self.department_id.name or ''}"
        t.cell(0,2).text = f"Emp. No. "
        
        t.cell(1, 0).merge(t.cell(1, 1))
        t.cell(1,0).text = f"DESIGNATION : {self.destination or ''}"
        t.cell(1,2).text = f"Date Of Joining :"
        
        # Title
        doc.add_heading("EMPLOYEE DATA FORM\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        t4 = doc.add_table(1, 3)
        t4.autofit = False

        t4.cell(0,0).text = ""
        
        cell = t4.rows[0].cells[1]
        p = cell.paragraphs[0]
        p.alignment = 1   # center
        run = p.add_run()
        
        # Get the absolute path to the image file
        module_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        image_path = os.path.join(module_path, "static", "images", "image.png")
        if os.path.exists(image_path):
            run.add_picture(image_path, width=Cm(3))
        
        t4.cell(0,2).text = f"Full Signature: ____________\nInitial: _____________________"
     

        # Personal Details
        ptext(f"1. Full Name : Mr./Ms. {self.candidate_id.partner_name or ''}",bold=True)
        ptext(f"    Father's Name : {self.father_name or ''}\n")


        # ========== table =========
        ptext("2. ADDRESS:", bold=True)
        t1 = table(5, 2)

        ctext(t1.cell(0,0), "PRESENT", bold=True)
        ctext(t1.cell(0,1), "PERMANENT", bold=True)

        ctext(t1.cell(1,0), self.present_address or "")
        ctext(t1.cell(1,1), self.permanent_address or "")

        ctext(t1.cell(2,0), f"Mobile: {str(self.partner_phone or '')}")
        ctext(t1.cell(2,1), f"Mobile: {str(self.mobile_no or '')}")

        ctext(t1.cell(3,0), "Telephone:")
        ctext(t1.cell(3,1), "Telephone:")

        t1.cell(4,0).merge(t1.cell(4,1))
        ctext(t1.cell(4,0), f"Email ID: {str(self.email_from or '')}")
        
        # Emergency Contact
        ptext("\n3. PERSON TO BE NOTIFIED IN CASE OF EMERGENCY:", bold=True)
        ptext(f"     Name: {self.emergency_contact_name or ''}")
        ptext("     Address: AS ABOVE")
        ptext(f"     Telephone: {str(self.emergency_contact_no or '')}\n")
                
                
        # Personal information 
        ptext("4. PERSONAL PARTICULARS:", bold=True)
        ptext(f"    a) Date of Birth : {str(self.dob or '')}")
        ptext(f"     b) Age : {self.age}")
        ptext(f"     c) Nationality : Indian")
        ptext(f"     d) Religion / Caste : HINDU")
        ptext(f"4. MARITAL STATUS : {str(self.marital_status or '')}\n", bold=True)


        # Family Details
        ptext("\n5. FAMILY PARTICULARS", bold=True)

        t2 = table(7, 5)
        ctext(t2.cell(0,0), "Sr No", bold=True)
        ctext(t2.cell(0,1), "Name", bold=True)
        ctext(t2.cell(0,2), "Relation", bold=True)
        ctext(t2.cell(0,3), "Age", bold=True)
        ctext(t2.cell(0,4), "Occupation", bold=True)


        family_data = [
            ("1", self.father_name or "", "Father"),
            ("2", self.mother_name or "", "Mother"),
            ("3", self.spouse_name or "", "Husband / Wife"),
        ]

        row = 1
        for sr, name, rel in family_data:
            ctext(t2.cell(row,0), sr)
            ctext(t2.cell(row,1), name)
            ctext(t2.cell(row,2), rel)
            row += 1
        
        # Educational Qualification
        ptext("\n6. EDUCATIONAL QUALIFICATION:", bold=True)

        t3 = table(len(self.education_ids) + 1, 6)
        
        ctext(t3.cell(0,0), "Sr No",bold=True)
        ctext(t3.cell(0,1), "Year",bold=True)
        ctext(t3.cell(0,2), "Exam / Course passed",bold=True)
        ctext(t3.cell(0,3), "School / College / University",bold=True)
        ctext(t3.cell(0,4), "Marks / Class",bold=True)
        ctext(t3.cell(0,5), "Remarks as regard to Scholarship / Distinction etc.",bold=True)
        
         
        for idx, rec in enumerate(self.education_ids, start=1):
            ctext(t3.cell(idx, 0), str(idx or ""))
            ctext(t3.cell(idx, 1), str(rec.year_of_passing or ""))
            ctext(t3.cell(idx, 2), str(rec.exam_passed or ""))
            ctext(t3.cell(idx, 3), str(rec.institution or ""))
            ctext(t3.cell(idx, 4), str(rec.marks_obtained or ""))
            ctext(t3.cell(idx, 5), "")
                    
            
        # Previous Employee Details
        ptext("\n7. PREVIOUS EMPLOYEMENT DETAILS:", bold=True)
        
        t4 = table(len(self.employment_ids) + 1, 7)

        
        ctext(t4.cell(0,0), "Sr No",bold=True)
        ctext(t4.cell(0,1), "Employer Name",bold=True)
        ctext(t4.cell(0,2), "Company",bold=True)
        ctext(t4.cell(0,3), "Address",bold=True)
        ctext(t4.cell(0,4), "Designation",bold=True)
        ctext(t4.cell(0,5), "Salary",bold=True)
        ctext(t4.cell(0,6), "Work Duration",bold=True)


        for idx, rec in enumerate(self.employment_ids, start=1):
            ctext(t4.cell(idx, 0), str(idx or ""))
            ctext(t4.cell(idx, 1), rec.employment_type_id.name or "")
            ctext(t4.cell(idx, 2), rec.company or "")
            ctext(t4.cell(idx, 3), rec.location or "")
            ctext(t4.cell(idx, 4), rec.designation or "")
            ctext(t4.cell(idx, 5), "")
            ctext(t4.cell(idx, 6), str(rec.years_of_experience or ""))

            
        # Declaration
        ptext("\n\nI HERE BY DECLARE AND CERTIFY THAT THE PARTICULARS / INFORMATION STATED IN THE APPLICATION IS TRUE / COMPLETE.")
        ptext("Signature : _______________________",align=WD_ALIGN_PARAGRAPH.RIGHT)
        
        
        # Save to memory
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Create attachment
        attachment = self.env['ir.attachment'].create({
            'name': f"Application_of_{self.candidate_id.partner_name or ''}.docx",
            'type': 'binary',
            'datas': base64.b64encode(buffer.getvalue()),
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'res_model': self._name,
            'res_id': self.id,
        })

        # Download
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }
        