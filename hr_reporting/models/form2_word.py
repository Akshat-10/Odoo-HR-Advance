from odoo import models
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from io import BytesIO
import base64
from odoo.tools.misc import format_date



def apply_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(ns.qn('w:val'), 'single')
        border.set(ns.qn('w:sz'), '8')
        border.set(ns.qn('w:space'), '0')
        border.set(ns.qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)


class HrCustomFormTwo(models.Model):
    _inherit = "hr.custom.form.form2"

    def action_download_form2_word(self):
        self.ensure_one()
        doc = Document()

        # ---------------- HEADER ----------------
        p = doc.add_heading("FORM 2 (Revised)", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph(
            "NOMINATION AND DECLARATION FORM FOR UNEXEMPTED /\n EXEMPTED ESTABLISHMENTS\n\n"
            
            "Declaration and Nomination Form under the Employees’ Provident Funds and\n"
            
            "Employees’ Pension Scheme\n\n"
            
            "(Paragraphs 33 & 61 (1) of the Employees Provident Fund Scheme, 1952 and\n"
            "Paragraph 18 of the Employees’ Pension Scheme, 1995)"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ---------------- BASIC DETAILS ----------------
        doc.add_paragraph(f"1. Name (in Block letters): {self.employee_id.name or ''}")
        doc.add_paragraph(f"2. Father’s / Husband’s Name: {self.father_husband_name or ''}")
        doc.add_paragraph(
            f"3. Date of Birth: {format_date(self.env, self.date_of_birth) if self.date_of_birth else ''}"
        )
        doc.add_paragraph(f"4. Sex: {self.gender or ''}")
        doc.add_paragraph(f"5. Marital Status: {self.marital_status or ''}")
        doc.add_paragraph(f"6. Account No.: {self.account_number or ''}")
        doc.add_paragraph(f"7. Address Permanent: {self.permanent_address or ''}")
        doc.add_paragraph(f"   Temporary: {self.temporary_address or ''}")
        doc.add_paragraph(
            f"8. Date of Joining: {format_date(self.env, self.date_of_joining) if self.date_of_joining else ''}"
        )

        # ---------------- PART A (EPF) ----------------
        p = doc.add_paragraph("\nPART – A (EPF)")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph(
            "I hereby nominate the person(s) mentioned below to receive the amount "
            "standing to my credit in the Employees’ Provident Fund in the event of my death:"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table_a = doc.add_table(rows=1, cols=6)
        hdr_cells = table_a.rows[0].cells
        titles = [
            "S.No", "Name & Address of Nominee", "Relationship",
            "Date of Birth", "Share (%)", "Guardian (if minor)"
        ]
        for i, title in enumerate(titles):
            hdr_cells[i].text = title
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, line in enumerate(self.part_a_line_ids, start=1):
            row = table_a.add_row().cells
            row[0].text = str(idx)
            row[1].text = f"{line.nominee_name or ''}\n{line.nominee_address or ''}"
            row[2].text = line.relationship or ""
            row[3].text = format_date(self.env, line.date_of_birth) if line.date_of_birth else ""
            row[4].text = line.share_amount or ""
            row[5].text = line.guardian_details or ""

        apply_table_borders(table_a)

        p = doc.add_paragraph(
            "* Certified that I have no family as defined in para 2(g) of EPF Scheme, 1952.\n"
            "* Certified that my father/mother is/are dependent upon me."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ---------------- SIGNATURE ----------------
        doc.add_paragraph("\nDate:")

        p = doc.add_paragraph("Signature or thumb impression of the subscriber")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("*Strike out whichever is not applicable.")

        # ---------------- PART B (EPS) ----------------
        p = doc.add_paragraph("\nPART – B (EPS)")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph(
            "I hereby furnish below particulars of the members of my family who would be "
            "eligible to receive widow/children pension in the event of my death."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table_b1 = doc.add_table(rows=1, cols=5)
        hdr = table_b1.rows[0].cells
        headers = ["S.No", "Name of Family Member", "Address", "Date of Birth", "Relationship"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, line in enumerate(self.part_b_family_line_ids, start=1):
            row = table_b1.add_row().cells
            row[0].text = str(idx)
            row[1].text = line.member_name or ""
            row[2].text = line.address or ""
            row[3].text = format_date(self.env, line.date_of_birth) if line.date_of_birth else ""
            row[4].text = line.relationship or ""

        apply_table_borders(table_b1)

        # ---------------- PART B NOMINEE ----------------
        doc.add_paragraph(
            "\nI hereby nominate the following person for receiving the monthly widow pension "
            "in the event of my death without leaving any eligible family member:"
        )

        table_b2 = doc.add_table(rows=1, cols=3)
        hdr = table_b2.rows[0].cells
        titles = ["Name & Address of Nominee", "Date of Birth", "Relationship"]
        for i, t in enumerate(titles):
            hdr[i].text = t
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in self.part_b_nominee_line_ids:
            row = table_b2.add_row().cells
            row[0].text = f"{line.nominee_name or ''}\n{line.nominee_address or ''}"
            row[1].text = format_date(self.env, line.date_of_birth) if line.date_of_birth else ""
            row[2].text = line.relationship or ""

        apply_table_borders(table_b2)

        # ---------------- SIGNATURE ----------------
        doc.add_paragraph("\nDate:")

        p = doc.add_paragraph("Signature or thumb impression of the subscriber")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("*Strike out whichever is not applicable.")

        # ================= CERTIFICATE BY EMPLOYER =================
        p = doc.add_paragraph("\nCERTIFICATE BY EMPLOYER")
        p.runs[0].bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            "Certified that the above declaration and nomination has been signed / thumb "
            "impressed before me by Shri / Smt. / Kum. ________________________________ "
            "______________________ employed in my establishment after he/she has read the "
            "entries / entries have been read over to him/her by me and got confirmed by him/her."
        )

        doc.add_paragraph("\nPlace : ____________________")
        p = doc.add_paragraph(
            "Signature of the employer or other\nAuthorized Officers of the Establishment"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = doc.add_paragraph("Designation : ____________________")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph("Dated the : ____________________")

        p = doc.add_paragraph(
            "Name & Address of the Factory /\nEstablishment or Rubber Stamp Thereon"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ================= GUIDANCE SECTION =================
        doc.add_page_break()
        # MAIN HEADING
        p = doc.add_heading("GUIDANCE FOR FILLING THE FORM No - 2", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        # EPF TITLE
        p = doc.add_paragraph("Employee’s Provident Fund Scheme, 1952 :-")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        p = doc.add_paragraph("( E P F )")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        # PARA 33
        p = doc.add_paragraph(
            "Para 33 :- Declaration by persons already employed at the time of institution of the fund :-")
        p.runs[0].bold = True

        doc.add_paragraph(
            "Para 33 :- Declaration by persons already employed at the time of institution of the fund :-\n"
            "Every person who is required or entitled to become a member of the Fund shall be asked "
            "forthwith by his employer to furnish and shall, on such demand, furnish to him, for "
            "communication to the Commissioner, particulars concerning himself and his nominee "
            "required for the declaration form in Form 2. Such employer shall enter the particulars "
            "in the declaration form and obtain the signature or thumb impression of the person concerned."
        )

        doc.add_paragraph("\nPara 61 : Nomination")
        p.runs[0].bold = True

        doc.add_paragraph(
            "1. Each member shall make in his declaration in Form 2, a nomination conferring the right "
            "to receive the amount that may stand to his credit in the Fund in the event of his death "
            "before the amount standing to his credit has become payable, or where the amount has "
            "become payable before payment has been made.\n\n"
            "2. A member may in this nomination distribute the amount that may stand to his credit in "
            "the Fund amongst his nominees at his own discretion.\n\n"
            "3. If a member has a family at the time of making a nomination, the nomination shall be "
            "in favour of one or more persons belonging to his family. Any nomination made by such "
            "member in favour of a person not belonging to his family shall be invalid.\n\n"
            "Provided that a fresh nomination shall be made by the member on his marriage and any "
            "nomination made before such marriage shall be deemed to be invalid.\n\n"
            "4. If at the time of making a nomination the member has no family, the nomination may be "
            "in favour of any person or persons but if the member subsequently acquires a family, "
            "such nomination shall forthwith be deemed to be invalid and the member shall make a "
            "fresh nomination in favour of one or more persons belonging to his family.\n\n"
            "4A. Where the nomination is wholly or partly in favour of a minor, the member may appoint "
            "a major person of his family to be the guardian of the minor nominee.\n\n"
            "5. A nomination made under sub-paragraph (1) may at any time be modified by a member "
            "after giving a written notice in Form 2.\n\n"
            "6. A nomination or its modification shall take effect to the extent that it is valid on "
            "the date on which it is received by the Commissioner."
        )

        doc.add_paragraph("\nPara 2(g) : Family Means :-")
        p.runs[0].bold = True

        doc.add_paragraph(
            "(i) In the case of a male member, his wife, his children, whether married or unmarried, "
            "his dependent parents and his deceased son’s widow and children.\n\n"
            "(ii) In the case of a female member, her husband, her children, whether married or "
            "unmarried, her dependent parents, her husband’s dependent parents, her deceased son’s "
            "widow and children.\n\n"
            "Explanation :- If the child of a member has been legally adopted by another person, "
            "such a child shall be considered as excluded from the family."
        )

        # EPS TITLE
        p = doc.add_paragraph("\nEMPLOYEES PENSION SCHEME, 1995")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        p = doc.add_paragraph("( E P S )")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        doc.add_paragraph(
            "Para 18 : Particulars to be supplied by the employees already employed at the time of "
            "commencement of the Employees’ Pension Scheme.\n\n"
            "Every person who is entitled to become a member of the Employees’ Pension Fund shall "
            "furnish particulars concerning himself and his family in the prescribed form."
        )

        # -------- PARA 2(vii) UNDER EPS --------
        p = doc.add_paragraph("\nPara 2(vii) :- Family means :-")
        p.runs[0].bold = True

        doc.add_paragraph(
            "(i) Wife in the case of male member of the Employees’ Pension Fund;\n"
            "(ii) Husband in the case of a female member of the Employees’ Pension Fund; and\n"
            "(iii) Sons and daughters of a member of the Employees’ Pension Fund;"
        )

        p = doc.add_paragraph("\nExplanation –")
        p.runs[0].bold = True

        doc.add_paragraph(
            "The expression “Sons” and “daughters” shall include children "
            "[legally adopted by the member]."
        )

        doc.add_paragraph(
            "\nNOTE : Members can nominate a person to receive benefits under the Employees’ Pension "
            "Scheme, 1995 where a member is unmarried or does not have any family."
        )

        # ---------------- SAVE & DOWNLOAD ----------------
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        employee_name = (
            self.employee_id.name.strip().replace(" ", "_")
            if self.employee_id and self.employee_id.name
            else "Employee"
        )

        attachment = self.env["ir.attachment"].create({
            "name": f"Form_2_{employee_name}.docx",
            "type": "binary",
            "datas": base64.b64encode(buffer.read()),
            "res_model": self._name,
            "res_id": self.id,
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        })

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }

