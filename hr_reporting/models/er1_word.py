from odoo import models
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from io import BytesIO
import base64


def apply_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(ns.qn('w:val'), 'single')
        border.set(ns.qn('w:sz'), '8')
        border.set(ns.qn('w:space'), '0')
        border.set(ns.qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)


class HrCustomFormER1(models.Model):
    _inherit = "hr.custom.form.er1"

    def action_download_er1_word(self):
        self.ensure_one()
        doc = Document()

        # ================= HEADER =================
        p = doc.add_heading("FORM ER-I", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph(
            "Quarterly return to be submitted to the local Employment Exchange for the Quarter ended\n"
            "The following information is required under the Employment Exchanges\n "
            "(Compulsory Notification of Vacancies) Act, 1959."
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ================= BASIC INFO (ALIGNED – NO BORDERS) =================
        info_table = doc.add_table(rows=6, cols=2)
        info_table.autofit = True

        info_data = [
            ("Quarter Ended", getattr(self, 'quarter_ended', '') or ''),
            ("Company", self.company_id.name if self.company_id else ''),
            ("Name of Employer", getattr(self, 'employer_name', '') or ''),
            ("Address of Employer",
             (getattr(self, 'employer_address', '') or '').replace('\n', ', ').replace('India', '').strip()),
            ("Whether Head Office / Branch Office", getattr(self, 'office_type', '') or ''),
            ("Nature of Business / Principal Activity",
             getattr(self, 'business_nature', '') or getattr(self, 'nature_of_business', '') or ''),
        ]

        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = f": {value}"

            info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
            info_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            info_table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ================= 1(a) EMPLOYMENT =================
        p = doc.add_paragraph("1(a) Employment Position")
        p.runs[0].bold = True

        doc.add_paragraph(
            "Total number of persons including working proprietors/partners/commission agents/"
            "contingent paid and contractual workers on the pay rolls of the establishment "
            "excluding part-time workers and apprentices."
            "(The figure should include every person whose wage or salary is paid by the establishment)"
        )

        table_emp = doc.add_table(rows=1, cols=3)
        hdr = table_emp.rows[0].cells
        hdr[0].text = "Man Power"
        hdr[1].text = "Previous Quarter (Last Working Day)"
        hdr[2].text = "Reporting Quarter (Last Working Day)"

        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in self.employment_line_ids:
            row = table_emp.add_row().cells
            row[0].text = line.manpower_name or ""
            row[1].text = str(line.prev_quarter_count or 0)
            row[2].text = str(line.current_quarter_count or 0)

        apply_table_borders(table_emp)

        doc.add_paragraph(
            "(b) Please indicate the main reasons for any increase or decrease in employment "
            "if the increase or decrease is more than 5% during the quarter."
        )
        doc.add_paragraph("..................................")
        doc.add_paragraph(
            "\nNote: Establishments are reminded of their obligation under the Employment Exchanges (Compulsory Notification of Vacancies)\nRules, 1960 for notifying the Employment Exchanges details of vacancies specified under the Act, before they are filled."
        )

        # ================= SECTION 2(a) =================
        p = doc.add_paragraph("\n2(a) Vacancies")
        p.runs[0].bold = True

        doc.add_paragraph(
            "Vacancies carrying total emoluments of Rs.60 or over per month, and of over 3 months duration."
        )
        doc.add_paragraph(
            "2(a) Number of vacancies occurred and notified during the quarter and the number filled during the quarter."
        )

        doc.add_paragraph("No. of vacancies which come within the purview of the Act.")

        table_vac = doc.add_table(rows=1, cols=5)
        hdr = table_vac.rows[0].cells
        titles = [
            "Occurred",
            "Notified to Local Employment Exchange",
            "Notified to Central Employment Exchange",
            "Filled",
            "Source (from which filled)",
        ]
        for i, t in enumerate(titles):
            hdr[i].text = t
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in self.vacancy_line_ids:
            row = table_vac.add_row().cells
            row[0].text = str(line.occurred or 0)
            row[1].text = str(line.notified_local or 0)
            row[2].text = str(line.notified_central or 0)
            row[3].text = str(line.filled_count or 0)
            row[4].text = line.source or ""

        apply_table_borders(table_vac)

        doc.add_paragraph(
            "\n2(b) Reasons for not notifying all vacancies occurred during the quarter under report view "
        )

        # ================= SECTION 3 =================
        p = doc.add_paragraph("\n3. Manpower Shortages, if any")
        p.runs[0].bold = True

        doc.add_paragraph(
            "Vacancies / Posts unfilled because of shortage of suitable applicants:"
        )

        table_short = doc.add_table(rows=1, cols=4)
        hdr = table_short.rows[0].cells
        headers = [
            "Occupation / Designation",
            "Essential Qualifications Prescribed",
            "Essential Experience Prescribed",
            "Experience Not Necessary",
        ]
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in self.shortage_line_ids:
            row = table_short.add_row().cells
            row[0].text = line.occupation_name or ""
            row[1].text = line.essential_qualification or ""
            row[2].text = line.essential_experience or ""
            row[3].text = line.experience_not_required or ""

        apply_table_borders(table_short)

        # ================= FOOTER =================
        doc.add_paragraph("\nPlace : ____________________")
        doc.add_paragraph("Date  : ____________________")

        p = doc.add_paragraph(self.company_id.name or "")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].bold = True

        p = doc.add_paragraph("Authorised Signatory")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph(
            "\nTo,\nThe Employment Exchange\n(Write address of Local Employment Exchange)"
        )

        doc.add_paragraph(
            "\nNote: This return shall relate to quarters ending 31st March / 30th June / "
            "30th September / 31st December and shall be rendered within 30 days after the end of the quarter concerned."
        )

        # ================= SAVE & DOWNLOAD =================
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # ========= FILE NAME LOGIC =========
        employer_name = (
            getattr(self, 'employer_name', '').strip().replace(" ", "_")
            if getattr(self, 'employer_name', '')
            else "Employer"
        )

        file_name = f"ER_1_{employer_name}.docx"

        attachment = self.env["ir.attachment"].create({
            "name": file_name,
            "type": "binary",
            "datas": base64.b64encode(buffer.read()),
            "res_model": self._name,
            "res_id": self.id,
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        })

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }
