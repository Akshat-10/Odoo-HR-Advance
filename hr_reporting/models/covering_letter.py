from odoo import models
from odoo import models
from odoo.http import request
from docx import Document
from docx.shared import Inches
import base64
from io import BytesIO
import tempfile
import os

class HrCustomFormCoverLetter(models.Model):
    _inherit = "hr.custom.form.cover_letter"

    def action_generate_covering_letter_report(self):
        return self.env.ref("hr_reporting.covering_letter_report").report_action(self, config=False)

    def action_download_covering_letter_word(self):
        self.ensure_one()

        doc = Document()

        # Logo
        if self.company_id and self.company_id.logo:
            image_stream = BytesIO(base64.b64decode(self.company_id.logo))
            doc.add_picture(image_stream, width=Inches(1.5))

        # Company
        doc.add_paragraph(self.company_id.name or "")
        doc.add_paragraph(self.company_id.partner_id.country_id.name or "")
        doc.add_paragraph("")

        # Date
        formatted_date = self.er_date.strftime("%d/%m/%Y") if self.er_date else ""
        doc.add_paragraph(f"Date: {formatted_date}")
        doc.add_paragraph("")

        # To
        doc.add_paragraph("To,")
        doc.add_paragraph(self.er_to_address or "")
        doc.add_paragraph("")

        # Subject
        p = doc.add_paragraph()
        p.add_run(f"Sub: {self.er_subject or ''}").bold = True
        doc.add_paragraph("")

        # Body
        doc.add_paragraph("Dear Sir / Madam,")
        doc.add_paragraph(self.er_body or "")
        doc.add_paragraph("")

        # Signature
        doc.add_paragraph("Thanking you,")
        doc.add_paragraph("Yours sincerely,")
        doc.add_paragraph("")
        doc.add_paragraph(f"({self.employee_id.name if self.employee_id else ''})")

        # Save to memory
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # ========= DOCUMENT REFERENCE BASED FILE NAME =========
        document_reference = (
            self.name.strip()
            .replace(" ", "_")
            .replace("/", "_")
            if self.name
            else "Document_Reference"
        )

        # Create attachment
        attachment = self.env["ir.attachment"].create({
            "name": f"Covering_Letter_{document_reference}.docx",
            "type": "binary",
            "datas": base64.b64encode(buffer.read()),
            "res_model": self._name,
            "res_id": self.id,
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        })

        # Download action
        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }




