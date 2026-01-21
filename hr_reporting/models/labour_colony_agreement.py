# -*- coding: utf-8 -*-
from odoo import models
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import base64
import re


class HrCustomFormLabourColonyAgreement(models.Model):
    _inherit = "hr.custom.form.labour_colony_agreement"

    def _strip_html_tags(self, html_content):
        """Convert HTML content to plain text, preserving list structure."""
        if not html_content:
            return ""
        
        # Replace list items with numbered format
        text = html_content
        
        # Remove HTML tags but preserve text content
        text = re.sub(r'<br\s*/?>', '\n', text)
        text = re.sub(r'</p>', '\n', text)
        text = re.sub(r'</li>', '\n', text)
        text = re.sub(r'<[^>]+>', '', text)
        
        # Clean up whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        
        return text

    def _get_policy_points(self, html_content):
        """Extract policy points from HTML content."""
        if not html_content:
            return []
        
        # Extract list items from HTML
        li_pattern = r'<li[^>]*>(.*?)</li>'
        matches = re.findall(li_pattern, html_content, re.DOTALL | re.IGNORECASE)
        
        points = []
        for match in matches:
            # Clean HTML tags from each point
            clean_text = re.sub(r'<[^>]+>', '', match).strip()
            if clean_text:
                points.append(clean_text)
        
        return points

    def _get_intro_paragraph(self, html_content):
        """Extract the introductory paragraph before the list."""
        if not html_content:
            return ""
        
        # Get text before <ol> or <ul>
        parts = re.split(r'<[ou]l[^>]*>', html_content, maxsplit=1)
        if parts:
            intro = re.sub(r'<[^>]+>', '', parts[0]).strip()
            return intro
        return ""

    def action_download_labour_colony_word(self):
        """Generate Word document for Labour Colony Agreement."""
        self.ensure_one()

        doc = Document()

        # ================= SET DEFAULT FONT =================
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # ================= COMPANY HEADER =================
        # Company Name
        company_name = self.company_id.name or "COMPANY NAME"
        p = doc.add_paragraph()
        run = p.add_run(company_name.upper())
        run.bold = True
        run.font.size = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Location
        location = self.location or ""
        if location:
            p = doc.add_paragraph()
            run = p.add_run(location.upper())
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # ================= ENGLISH TITLE =================
        p = doc.add_paragraph()
        run = p.add_run("Policy for Accommodation / Residence provided by the Company")
        run.bold = True
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ================= HINDI TITLE =================
        hindi_title = self.agreement_content_hindi_title or "कंपनीद्वाराप्रदानकीगईआवास / निवासकेलिएनीति"
        p = doc.add_paragraph()
        run = p.add_run(hindi_title)
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # ================= INTRO PARAGRAPH =================
        intro_text = self._get_intro_paragraph(self.agreement_content)
        if intro_text:
            p = doc.add_paragraph()
            run = p.add_run(intro_text)
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph()  # Spacing

        # ================= POLICY POINTS =================
        policy_points = self._get_policy_points(self.agreement_content)
        
        for idx, point in enumerate(policy_points, start=1):
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}.\t{point}")
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Add spacing between points
            p.paragraph_format.space_after = Pt(6)

        # ================= ADD SPACING BEFORE SIGNATURES =================
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        # ================= SIGNATURE SECTION =================
        # Create a table for signatures (1 row, 2 columns)
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_table.autofit = True

        # Left signature - SIGN OF AUTHORITY
        left_cell = sig_table.cell(0, 0)
        left_p = left_cell.paragraphs[0]
        run = left_p.add_run("SIGN OF AUTHORITY")
        run.bold = True
        run.font.size = Pt(11)
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Right signature - SIGN OF EMPLOYEE
        right_cell = sig_table.cell(0, 1)
        right_p = right_cell.paragraphs[0]
        run = right_p.add_run("SIGN OF EMPLOYEE")
        run.bold = True
        run.font.size = Pt(11)
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ================= SAVE DOCUMENT =================
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Document reference for filename
        document_reference = (
            self.name.strip()
            .replace(" ", "_")
            .replace("/", "_")
            if self.name
            else "Labour_Colony_Agreement"
        )

        employee_name = (
            self.employee_id.name.strip()
            .replace(" ", "_")
            if self.employee_id and self.employee_id.name
            else ""
        )

        filename = f"Labour_Colony_Agreement_{document_reference}"
        if employee_name:
            filename += f"_{employee_name}"
        filename += ".docx"

        # Create attachment
        attachment = self.env["ir.attachment"].create({
            "name": filename,
            "type": "binary",
            "datas": base64.b64encode(buffer.read()),
            "res_model": self._name,
            "res_id": self.id,
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        })

        # Return download action
        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }
