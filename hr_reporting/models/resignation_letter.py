from odoo import models
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

class HrCustomFormResignationLetter(models.Model):
    _inherit = "hr.custom.form.resignation_letter"
    
    def action_generate_resignation_report(self):
        return self.env.ref('hr_reporting.resignation_letter_report').report_action(self,config=False) 


    def action_download_word(self):
        self.ensure_one()

        doc = Document()

        # Common font style
        style = doc.styles['Normal']
        style.font.name = 'Calibri (Body)'
        style.font.size = Pt(11)

        # Date
        p = doc.add_paragraph()
        p.add_run(f"Date. {self.resignation_date.strftime('%d/%m/%Y') or ''}\n")

        # Address block
        doc.add_paragraph(
            "To,\n"
            "The General Manager,\n"
            "GERMAN TMT PVT LTD,\n" 
        )

        # Subject & Department
        doc.add_paragraph(
            f"Sub: Resignation from the post of {self.employee_id.job_id.name or '____________________'}\n"
            f"DEPARTMENT {self.employee_id.department_id.name or '____________________'}"
        )

        # Salutation
        doc.add_paragraph("Dear\nSir,\n")

        # Main body
        doc.add_paragraph(
            f"I the undersigned Ms/Mr. {self.employee_id.name or '____________________'} "
            f"S/O {self.employee_id.father_name or '____________________'} is working "
            f"{self.employee_id.job_id.name or '____________________'} as with your Company. "
            "Due to my personal reason salary not affordable to me "
            "could not work further more with Company and hence "
            f"tendering my Resignation from the post of "
            f"{self.employee_id.job_id.name or '____________________'}."
        )

        doc.add_paragraph(
            "I hereby declare that I have received all my dues from company "
            "and there are no dues by what so ever name called pending to be "
            "received from your company and that I shall not claim any amount "
            "from the company in future. My account has been cleared and settled."
        )

        doc.add_paragraph(
            "You are requested to accept my resignation with immediate effect "
            "and relieve me from the services. I am thankful for giving all the "
            "support and guidance during my stay."
        )

        # Closing
        doc.add_paragraph(
            "\nThanking you,\n"
            "Your's truly\n\n"
            f"({self.employee_id.name or '____________________'})"
        )

        p = doc.add_paragraph(
            "\nApproved By,\n\n"
            "____________________\n"
            "Authorized Signatory\n"
            "GERMAN TMT PVT LTD"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to memory
        buffer = io.BytesIO()
        doc.save(buffer)
        
        filename = (
            f"Resignation Letter - {self.employee_id.name or ''}({self.employee_id.employee_code or ''}).docx"
            if self.employee_id.employee_code
            else f"Resignation Letter - {self.employee_id.name or ''}.docx"
        )

        # Create attachment
        attachment = self.env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(buffer.getvalue()),
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'res_model': self._name,
            'res_id': self.id,
        })

        # Download
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }