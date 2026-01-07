from odoo import models, fields, api
from io import BytesIO
import base64


class HrCustomFormNominationF(models.Model):
    _inherit = "hr.custom.form.nomination_f"

    def action_generate_nomination_f_docx(self):
        self.ensure_one()

        # ================= SAFE IMPORTS =================
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from bs4 import BeautifulSoup

        document = Document()


        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        company = self.company_id or self.env.company

        # ================= HEADER =================
        p = document.add_paragraph("Form ‘F’")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        p = document.add_paragraph("[See sub-rule (1) of rule 6]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph("\nNomination\n")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

        # ================= TO – COMPANY =================
        document.add_paragraph("To")
        document.add_paragraph(
            f"{company.name}\n"
            f"{company.street or ''}\n"
            f"{company.city or ''} - {company.zip or ''}"
        )

        document.add_paragraph("")

        # ==================================================
        # 1️⃣ NOMINATION CONTEXT (FROM TAB)
        # ==================================================
        if self.nomination_context:
            soup = BeautifulSoup(self.nomination_context, "html.parser")
            for element in soup.find_all(["p", "li"]):
                text = element.get_text(strip=True)
                if text:
                    para = document.add_paragraph(text)
                    para.paragraph_format.space_after = Pt(6)

        # ==================================================
        # 2️⃣ NOMINEES (FROM NOMINEES TAB)
        # ==================================================
        document.add_paragraph("\nNominee(s):")

        nominee_table = document.add_table(rows=1, cols=4)
        nominee_table.style = "Table Grid"

        hdr = nominee_table.rows[0].cells
        hdr[0].text = "Name in full with full address of nominee(s)"
        hdr[1].text = "Relationship with the employee"
        hdr[2].text = "Age of nominee"
        hdr[3].text = "Proportion by which the gratuity will be shared"

        for line in self.nominee_line_ids.sorted("sequence"):
            row = nominee_table.add_row().cells
            row[0].text = line.nominee_name_address or ""
            row[1].text = line.relationship_id.name if line.relationship_id else ""
            row[2].text = line.age or ""
            row[3].text = f"{line.share_percentage or 0:.2f}%"

        # ==================================================
        # 3️⃣ STATEMENT (FROM STATEMENT TAB)
        # ==================================================
        document.add_paragraph("\nStatement")

        document.add_paragraph(f"1. Name of employee in full : {self.employee_id.name or ''}")
        document.add_paragraph(
            f"2. Sex : {dict(self._fields['gender'].selection).get(self.gender, '')}"
        )
        document.add_paragraph(
            f"3. Religion : {self.caste_id.name if self.caste_id else ''}"
        )
        document.add_paragraph(
            f"4. Whether unmarried/married/widow/widower : "
            f"{dict(self._fields['marital_status'].selection).get(self.marital_status, '')}"
        )
        document.add_paragraph(
            f"5. Department/Branch/Section where employed : "
            f"{self.employee_id.department_id.name if self.employee_id.department_id else ''}"
        )
        document.add_paragraph(
            f"6. Post held with Ticket or Serial No., if any : {self.post_held or ''}"
        )

        document.add_paragraph(f"\nJoining Date : {self.employee_id.first_contract_date or ''}")
        document.add_paragraph(f"Permanent Address : {self.permanent_address or ''}")

        # ================= LOCATION DETAILS =================
        document.add_paragraph("\nLocation Details")
        document.add_paragraph(f"Village : {self.village or ''}")
        document.add_paragraph(f"Thana : {self.thana or ''}")
        document.add_paragraph(f"Sub-division : {self.sub_division or ''}")
        document.add_paragraph(f"Post Office : {self.post_office or ''}")
        document.add_paragraph(f"District : {self.district or ''}")
        document.add_paragraph(f"State : {self.state or ''}")

        # ==================================================
        # 4️⃣ WITNESSES (FROM WITNESSES TAB)
        # ==================================================
        document.add_paragraph("\nDeclaration by witnesses")

        for idx, witness in enumerate(self.witness_line_ids.sorted("sequence"), start=1):
            document.add_paragraph(
                f"{idx}. {witness.witness_name}\n"
                f"   {witness.witness_address}"
            )

        # ================= SAVE WORD FILE =================
        output = BytesIO()
        document.save(output)
        output.seek(0)

        attachment = self.env["ir.attachment"].create({
            "name": "Nomination_Form_F.docx",
            "type": "binary",
            "datas": base64.b64encode(output.read()),
            "res_model": self._name,
            "res_id": self.id,
        })

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }
