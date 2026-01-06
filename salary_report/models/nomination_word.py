from odoo import models, fields, api
from io import BytesIO
import base64
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class HrCustomFormNominationF(models.Model):
    _inherit = "hr.custom.form.nomination_f"

    def action_generate_nomination_f_docx(self):
        self.ensure_one()

        # ================= SAFE IMPORTS =================
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from bs4 import BeautifulSoup

        # 🔴 CHANGE START: imports for cell background color
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        # 🔴 CHANGE END

        document = Document()

        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)

        company = self.company_id or self.env.company

        # 🔴 CHANGE START: Date format d/m/Y
        formatted_date = self.form_date.strftime("%d/%m/%Y") if self.form_date else ""
        # 🔴 CHANGE END

        # ================= HEADER =================
        p = document.add_paragraph("Form ‘F’")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        p = document.add_paragraph("[See sub-rule (1) of rule 6]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph("\nNomination\n")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        company = self.env.company
        partner = company.partner_id

        document.add_paragraph("To")

        p = document.add_paragraph()
        run = p.add_run(company.name)
        run.bold = True

        p.add_run(
            f"\n{partner.street or ''}"
            f"\n{(', ' + partner.street2) if partner.street2 else ''}\n"
            f"{partner.city or ''} - {partner.zip or ''}"
        )

        document.add_paragraph("")

        # ==================================================
        # 1️⃣ NOMINATION CONTEXT
        # ==================================================
        if self.nomination_context:
            soup = BeautifulSoup(self.nomination_context, "html.parser")
            for element in soup.find_all(["p", "li"]):
                text = element.get_text(strip=True)
                if text:
                    para = document.add_paragraph(text)
                    para.paragraph_format.space_after = Pt(6)

        # ==================================================
        # 2️⃣ NOMINEES
        # ==================================================
        p = document.add_paragraph("\nNominee(s):")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        nominee_table = document.add_table(rows=1, cols=4)
        nominee_table.style = "Table Grid"
        nominee_table.allow_autofit = False
        nominee_table.columns[0].width = Inches(3.5)
        nominee_table.columns[1].width = Inches(1.8)
        nominee_table.columns[2].width = Inches(1.2)
        nominee_table.columns[3].width = Inches(2.0)

        hdr = nominee_table.rows[0].cells
        hdr[0].text = "Name in full with full address of nominee(s)"
        hdr[1].text = "Relationship with the employee"
        hdr[2].text = "Age of nominee"
        hdr[3].text = "Proportion by which the gratuity will be shared"

        # 🔴 CHANGE START: helper to set cell background color
        def set_cell_bg(cell, color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), color)
            tcPr.append(shd)
        # 🔴 CHANGE END

        # 🔴 CHANGE START: Header row blue color + bold
        for cell in hdr:
            set_cell_bg(cell, "9DCBFF")
            for run in cell.paragraphs[0].runs:
                run.bold = True
        # 🔴 CHANGE END

        for line in self.nominee_line_ids.sorted("sequence"):
            row = nominee_table.add_row().cells
            row[0].text = line.nominee_name_address or ""
            row[1].text = line.relationship_id.name if line.relationship_id else ""
            row[2].text = line.age or ""
            row[3].text = f"{line.share_percentage or 0:.2f}%"

            # 🔴 CHANGE START: light blue data row color
            for cell in row:
                set_cell_bg(cell, "E7F3FF")
            # 🔴 CHANGE END

        # ==================================================
        # 3️⃣ STATEMENT
        # ==================================================
        p = document.add_paragraph("\nStatement")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        document.add_paragraph(f"1. Name of employee in full : {self.employee_id.name or ''}")
        document.add_paragraph(f"2. Sex : {dict(self._fields['gender'].selection).get(self.gender, '')}")
        document.add_paragraph(f"3. Religion : {self.caste_id.name if self.caste_id else ''}")
        document.add_paragraph(
            f"4. Whether unmarried/married/widow/widower : "
            f"{dict(self._fields['marital_status'].selection).get(self.marital_status, '')}"
        )
        document.add_paragraph(
            f"5. Department/Branch/Section where employed : "
            f"{self.employee_id.department_id.name if self.employee_id.department_id else ''}"
        )
        document.add_paragraph(f"6. Post held with Ticket or Serial No., if any : {self.post_held or ''}")

        joining_date = (
            self.employee_id.first_contract_date.strftime("%d/%m/%Y")
            if self.employee_id.first_contract_date
            else ""
        )

        document.add_paragraph(f"7. Joining Date : {joining_date}")

        document.add_paragraph(f"8. Permanent Address : {self.permanent_address or ''}")

        # ================= LOCATION DETAILS =================
        p = document.add_paragraph("\nLocation Details")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(13)

        document.add_paragraph(f"Village : {self.village or ''}")
        document.add_paragraph(f"Thana : {self.thana or ''}")
        document.add_paragraph(f"Sub-division : {self.sub_division or ''}")
        document.add_paragraph(f"Post Office : {self.post_office or ''}")
        document.add_paragraph(f"District : {self.district or ''}")
        document.add_paragraph(f"State : {self.state or ''}")

        # ==================================================
        # DECLARATION BY WITNESSES
        # ==================================================
        p = document.add_paragraph("\nDeclaration by witnesses")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        p = document.add_paragraph("Nomination signed/thumb impressed before me.\n")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create a table with 2 columns for proper alignment
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False

        # Set column widths
        table.columns[0].width = Inches(3.5)  # Left column for name/address
        table.columns[1].width = Inches(3.0)  # Right column for signature

        # Remove table borders
        def remove_table_borders(table):
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            tblPr.append(tblBorders)

        remove_table_borders(table)

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Name in full and full address of witnesses."
        header_cells[1].text = "Signature of witnesses."

        # ================== WITNESSES ==================
        witness_number = 1
        for witness in self.witness_line_ids.sorted("sequence"):
            # Add new row
            row_cells = table.add_row().cells

            # Left cell: Number, Name, and Address
            left_cell = row_cells[0]
            p_left = left_cell.paragraphs[0]
            p_left.add_run(f"{witness_number}. {witness.witness_name}\n")
            p_left.add_run(f"     {witness.witness_address}")
            p_left.paragraph_format.space_after = Pt(12)

            # Right cell: Signature text only
            right_cell = row_cells[1]
            p_right = right_cell.paragraphs[0]
            p_right.text = "Signature :"
            p_right.paragraph_format.space_after = Pt(12)

            witness_number += 1
        # ==================================================
        # CERTIFICATE BY EMPLOYER
        # ==================================================
        p = document.add_paragraph("\nCertificate by the employer")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        document.add_paragraph(
            "Certified that the particulars of the above nomination have been verified "
            "and recorded in this establishment."
        )
        employee_code = self.employee_id.employee_code if self.employee_id and self.employee_id.employee_code else ""

        document.add_paragraph(
            f"Employer's Reference No. - {employee_code} FOR GTP\n"
        )

        company = self.env.company
        partner = company.partner_id

        document.add_paragraph("To")

        p = document.add_paragraph()
        run = p.add_run(company.name)
        run.bold = True

        p.add_run(
            f"\n{partner.street or ''}"
            f"\n{(', ' + partner.street2) if partner.street2 else ''}\n"
            f"{partner.city or ''} - {partner.zip or ''}"
        )
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph(f"\nDate : {formatted_date}")

        # ==================================================
        # ACKNOWLEDGEMENT BY EMPLOYEE
        # ==================================================
        p = document.add_paragraph("\nAcknowledgement by the employee")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(12)

        document.add_paragraph(
            "Received the duplicate copy of nomination in Form 'F' "
            "filed by me and duly certified by the employer."
        )

        sign_table = document.add_table(rows=1, cols=2)
        sign_table.allow_autofit = False
        sign_table.columns[0].width = Inches(3.5)
        sign_table.columns[1].width = Inches(3.5)

        left_cell = sign_table.rows[0].cells[0]
        right_cell = sign_table.rows[0].cells[1]

        p = left_cell.paragraphs[0]
        p.text = f"Date : {formatted_date}"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p = right_cell.paragraphs[0]
        p.text = "Signature of the employee"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ==================================================
        # SAVE FILE
        # ==================================================
        output = BytesIO()
        document.save(output)
        output.seek(0)

        attachment = self.env["ir.attachment"].create({
            "name": "Nomination_Form_F.docx",
            "type": "binary",
            "datas": base64.b64encode(output.read()),
            "res_model": self._name,
            "res_id": self.id,
        })

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/{attachment.id}?download=true",
            "target": "self",
        }
