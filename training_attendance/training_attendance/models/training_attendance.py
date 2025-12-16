from odoo import models, fields
from odoo.exceptions import UserError
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from io import BytesIO
import base64
import tempfile


class TrainingAttendance(models.Model):
    _name = "training.attendance"
    _description = "Training Attendance"

    program_name = fields.Char(string="Program Name",)
    department_id = fields.Many2one("hr.department", string="Department")
    faculty_name = fields.Char(string="Faculty Name")
    venue = fields.Char(string="Venue")
    date = fields.Date(string="Date")
    time = fields.Char(string="Time")

    training_given_to = fields.Selection([
        ("company", "Company Employee"),
        ("contractor", "Contractor Employee"),
    ], string="Training Given To")

    employee_line_ids = fields.One2many(
        "training.attendance.line",
        "attendance_id",
        string="Employees"
    )

    # existing fields remain same …

    word_file = fields.Binary(string="Training Attendance Word")
    word_filename = fields.Char(string="File Name")

    def action_download_word(self):
        self.ensure_one()

        doc = Document()

        # ================= PAGE OUTER BORDER =================
        section = doc.sections[0]
        sectPr = section._sectPr

        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(ns.qn('w:offsetFrom'), 'page')

        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(ns.qn('w:val'), 'single')
            border.set(ns.qn('w:sz'), '12')  # thickness
            border.set(ns.qn('w:space'), '24')  # distance from text
            border.set(ns.qn('w:color'), '000000')
            pgBorders.append(border)

        sectPr.append(pgBorders)

        # ================= HEADER WITH LOGO =================
        header_table = doc.add_table(rows=1, cols=3)
        header_table.autofit = False

        header_table.columns[0].width = Inches(1.5)  # Logo
        header_table.columns[1].width = Inches(4.5)  # Title center
        header_table.columns[2].width = Inches(1.5)  # Balance

        left_cell = header_table.cell(0, 0)
        center_cell = header_table.cell(0, 1)
        right_cell = header_table.cell(0, 2)

        # ---- LEFT: COMPANY LOGO ----
        company = self.env.company
        if company.logo:
            logo_bytes = base64.b64decode(company.logo)
            tmp_logo = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp_logo.write(logo_bytes)
            tmp_logo.close()

            p = left_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(tmp_logo.name, width=Inches(1.2))
        else:
            left_cell.text = ""

        # ---- CENTER: TITLE (SINGLE LINE) ----
        p_title = center_cell.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.keep_together = True

        run_title = p_title.add_run("TRAINING ATTENDANCE")
        run_title.bold = True
        run_title.font.size = Pt(14)

        # ---- COMPANY NAME (SINGLE LINE) ----
        p_company = center_cell.add_paragraph()
        p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_company.paragraph_format.keep_together = True

        run_company = p_company.add_run("GERMAN TMT PRIVATE LIMITED")
        run_company.bold = True
        run_company.font.size = Pt(16)

        right_cell.text = ""

        doc.add_paragraph("")

        # ================= TRAINING DETAILS (LINE STYLE) =================
        def add_label_line(label, value=""):
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label} : ")
            run_label.bold = True

            run_value = p.add_run(value if value else " " * 60)
            run_value.underline = True

        add_label_line("Program Name", self.program_name or "")
        add_label_line("Department", self.department_id.name if self.department_id else "")
        add_label_line("Faculty Name", self.faculty_name or "")
        add_label_line("Venue", self.venue or "")
        add_label_line("Date / Time", f"{self.date or ''}  {self.time or ''}")
        add_label_line(
            "Training Given To",
            dict(self._fields['training_given_to'].selection).get(self.training_given_to, "")
        )

        doc.add_paragraph("")

        # ================= EMPLOYEE TABLE =================
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Sr. No"
        hdr_cells[1].text = "Name of Employee"
        hdr_cells[2].text = "Emp Code"
        hdr_cells[3].text = "Department"
        hdr_cells[4].text = "Designation"
        hdr_cells[5].text = "Signature"

        lines = self.employee_line_ids
        total_rows = max(20, len(lines))

        for i in range(total_rows):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)

            if i < len(lines):
                line = lines[i]
                row_cells[1].text = line.employee_name or ""
                row_cells[2].text = line.employee_code or ""
                row_cells[3].text = line.department_id.name if line.department_id else ""
                row_cells[4].text = line.designation or ""
                row_cells[5].text = ""
            else:
                for col in range(1, 6):
                    row_cells[col].text = ""

        doc.add_paragraph("\nSafety Officer Signature: _______________________")

        # ================= SAVE FILE =================
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        self.word_file = base64.b64encode(buffer.read())
        self.word_filename = "Training_Attendance.docx"

        return {
            "type": "ir.actions.act_url",
            "url": (
                f"/web/content/?model=training.attendance&id={self.id}"
                f"&field=word_file&filename_field=word_filename&download=true"
            ),
            "target": "self",
        }


class TrainingAttendanceLine(models.Model):
    _name = "training.attendance.line"
    _description = "Training Attendance Line"

    attendance_id = fields.Many2one("training.attendance", ondelete="cascade")
    sr_no = fields.Integer(string="Sr. No")
    employee_name = fields.Char(string="Employee Name")
    employee_code = fields.Char(string="Employee Code")
    department_id = fields.Many2one("hr.department", string="Department")
    designation = fields.Char(string="Designation")
    signature = fields.Char(string="Signature")
    attachment = fields.Binary(string="Attachment")
    attachment_filename = fields.Char(string="File Name")



